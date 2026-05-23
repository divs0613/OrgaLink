import os
import sys
import subprocess

def run_pip(args):
    print(f"Running pip command: {args}")
    subprocess.check_call([sys.executable, "-m", "pip"] + args + ["--break-system-packages"])

if __name__ == "__main__":
    # Clean up old legacy docx package and install modern python-docx
    try:
        # Uninstall old legacy docx
        subprocess.run([sys.executable, "-m", "pip", "uninstall", "-y", "docx"], capture_output=True)
        
        # Install modern python-docx and reportlab
        run_pip(["install", "reportlab", "python-docx"])
    except Exception as e:
        print(f"Dependency setup failed: {e}")
        print("Please manually run: pip install reportlab python-docx")
        sys.exit(1)

    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from docx import Document

    # 1. Generate Sample Developer PDF
    pdf_path = "sample_developer.pdf"
    print(f"Generating {pdf_path}...")
    c = canvas.Canvas(pdf_path, pagesize=letter)
    c.drawString(100, 750, "Applicant Profile: Alex Rivera (Developer Match)")
    c.drawString(100, 720, "Student ID: 2023-10043")
    c.drawString(100, 690, "Bio:")
    c.drawString(100, 670, "I am a web application developer with strong skills in writing clean code.")
    c.drawString(100, 650, "I build responsive user interfaces using HTML, CSS, and modern JavaScript.")
    c.drawString(100, 630, "I am highly proficient in React, Node.js, Python, and SQL databases.")
    c.drawString(100, 610, "I use Git for collaborative development. Recommended for the Web Development committee.")
    c.save()

    # 2. Generate Sample Marketing DOCX
    docx_path = "sample_marketing.docx"
    print(f"Generating {docx_path}...")
    doc = Document()
    doc.add_heading("Applicant Profile: Taylor Chen (Marketing Match)", 0)
    doc.add_paragraph("Student ID: 2023-11202")
    doc.add_heading("Biography & Experience", level=1)
    doc.add_paragraph(
        "I am a creative digital designer specializing in marketing, social media campaigns, and brand identities. "
        "I have advanced skills in Photoshop, Illustrator, and Figma. I love UI/UX design, creating video graphics, "
        "and content creation. I focus on branding, copywriting, and search engine optimization (SEO) to drive audience engagement."
    )
    doc.save(docx_path)

    print("Success! Test samples 'sample_developer.pdf' and 'sample_marketing.docx' generated in workspace root.")
